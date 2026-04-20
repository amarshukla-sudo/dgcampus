"""
DG Sheet Generator  v10.0
Created by Dr. Amar Shukla · IILM University, Gurugram
=========================================================
FIXES in v10:
  1. Session sheet: ALL cells written as pure TEXT strings (no number-format issues)
     → matches exact DL.xlsx format — no more AWS Athena parse errors
  2. Attendance: CORRECT matching logic
     → template students not in master → kept PRESENT (not forced ABSENT)
     → only students whose master att==0 get marked ABSENT
  3. UI: Ask "how many sessions", "how many units to include"
  4. Date selection: per-month individual date checkboxes
  5. Works with ANY file format for attendance and syllabus
"""

import io, re, zipfile, warnings
from collections import defaultdict, OrderedDict
from datetime import datetime, date

import pandas as pd
import streamlit as st
from openpyxl import load_workbook, Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

warnings.filterwarnings("ignore")

# ──────────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="DG Sheet Generator — Dr. Amar Shukla",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown("""
<style>
html, body, [class*="css"] { font-size: 16px !important; }
.main { padding: 0.5rem 2rem 3rem; }

.banner {
    background: linear-gradient(135deg, #0a1f35 0%, #1565c0 100%);
    color: #fff; padding: 1.5rem 2.5rem; border-radius: 16px; margin-bottom: 1.2rem;
}
.banner h1 { margin: 0; font-size: 2rem; font-weight: 900; }
.banner .sub { font-size: 1rem; opacity: .88; margin: .35rem 0 0; }
.banner .credit { font-size: .88rem; opacity: .68; margin-top: .5rem; font-style: italic; }

.sec { font-size: 1.18rem; font-weight: 800; color: #0a1f35;
       border-left: 6px solid #1565c0; padding: .3rem 0 .3rem .75rem;
       margin: 1.2rem 0 .6rem; }

.al-info { background:#e3f2fd; border-left:5px solid #1565c0; border-radius:0 8px 8px 0;
           padding:.6rem 1.1rem; margin:.35rem 0; font-size:.95rem; color:#0d47a1; }
.al-ok   { background:#e8f5e9; border-left:5px solid #2e7d32; border-radius:0 8px 8px 0;
           padding:.6rem 1.1rem; margin:.35rem 0; font-size:.95rem; color:#1b5e20; }
.al-warn { background:#fff8e1; border-left:5px solid #f57f17; border-radius:0 8px 8px 0;
           padding:.6rem 1.1rem; margin:.35rem 0; font-size:.95rem; color:#e65100; }
.al-err  { background:#ffebee; border-left:5px solid #c62828; border-radius:0 8px 8px 0;
           padding:.6rem 1.1rem; margin:.35rem 0; font-size:.95rem; color:#b71c1c; }

.stat-row { display:flex; gap:.9rem; margin: .7rem 0; flex-wrap:wrap; }
.stat-card { background:#fff; border:2px solid #e3f2fd; border-radius:12px;
             padding:.9rem 1.2rem; text-align:center; flex:1; min-width:120px;
             box-shadow:0 2px 8px rgba(21,101,192,.09); }
.stat-num { font-size:2.2rem; font-weight:900; color:#1565c0; line-height:1.1; }
.stat-lbl { font-size:.83rem; color:#546e7a; margin-top:.25rem; font-weight:600; }

.date-chip { display:inline-block; background:#1565c0; color:#fff;
             border-radius:14px; padding:.15rem .55rem; margin:.1rem;
             font-size:.77rem; font-weight:700; }

.time-tag { font-size:.82rem; font-weight:700; background:#1565c0; color:#fff;
            border-radius:6px; padding:.2rem .5rem; margin-top:.3rem; display:inline-block; }

.stButton>button { font-size:1rem !important; font-weight:700 !important;
                   padding:.6rem 1.5rem !important; border-radius:10px !important; }
.stDownloadButton>button { background:#0a1f35 !important; color:#fff !important;
    border-radius:10px !important; font-weight:700 !important;
    font-size:.95rem !important; width:100% !important; padding:.55rem 1rem !important; }
.stDownloadButton>button:hover { background:#1565c0 !important; }

label { font-size:1rem !important; font-weight:600 !important; }
.streamlit-expanderHeader { font-size:1rem !important; font-weight:700 !important; }
.badge { display:inline-block; background:#e8f5e9; color:#2e7d32;
         padding:.15rem .65rem; border-radius:12px; font-size:.8rem; font-weight:700; }
hr { border:none; border-top:2.5px solid #e3f2fd; margin:1.1rem 0; }
.footer { text-align:center; color:#90a4ae; font-size:.82rem; margin-top:2rem;
          padding-top:1rem; border-top:2px solid #e3f2fd; }
</style>
""", unsafe_allow_html=True)

def al_info(m): st.markdown(f'<div class="al-info">ℹ️  {m}</div>',  unsafe_allow_html=True)
def al_ok(m):   st.markdown(f'<div class="al-ok">✅  {m}</div>',   unsafe_allow_html=True)
def al_warn(m): st.markdown(f'<div class="al-warn">⚠️  {m}</div>', unsafe_allow_html=True)
def al_err(m):  st.markdown(f'<div class="al-err">❌  {m}</div>',  unsafe_allow_html=True)
def sec(t):     st.markdown(f'<div class="sec">{t}</div>',          unsafe_allow_html=True)

ALL_DAYS = ["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday","Sunday"]
SHORT    = {d: d[:3] for d in ALL_DAYS}

# ══════════════════════════════════════════════════════════════════════
# UTILITIES
# ══════════════════════════════════════════════════════════════════════

def safe_isna(v) -> bool:
    try: return bool(pd.isna(v))
    except: return False

def parse_date(raw) -> date | None:
    if raw is None or safe_isna(raw): return None
    if isinstance(raw, datetime): return raw.date()
    if isinstance(raw, date):     return raw
    s = str(raw).strip()
    if not s or s.lower() in ("nat","nan","none","pd.nat"): return None
    s = re.sub(r"\([^)]*\)","",s); s = re.sub(r"\b(sub|Sub|SUB)\b","",s)
    s = re.sub(r"\s+"," ",s).strip()
    if not s: return None
    for fmt in ("%d %b %Y","%d %B %Y","%Y-%m-%d %H:%M:%S",
                "%Y-%m-%d","%d-%m-%Y","%d/%m/%Y","%d %b %y"):
        try: return datetime.strptime(s, fmt).date()
        except: pass
    try:
        r = pd.to_datetime(s, dayfirst=True, errors="coerce")
        return None if pd.isna(r) else r.date()
    except: return None

def to_24h(h: int, m: str, ampm: str) -> str:
    hh = int(h)
    if ampm == "PM" and hh != 12: hh += 12
    if ampm == "AM" and hh == 12: hh = 0
    return f"{hh:02d}:{m}"

def best_sheet(fb: bytes, hints: list = None) -> pd.DataFrame:
    xf = pd.ExcelFile(io.BytesIO(fb)); ns = xf.sheet_names
    if len(ns)==1: return pd.read_excel(io.BytesIO(fb), sheet_name=ns[0], header=None)
    if hints:
        for kw in hints:
            for s in ns:
                if kw.lower() in s.lower():
                    return pd.read_excel(io.BytesIO(fb), sheet_name=s, header=None)
    best, most = ns[0], -1
    for s in ns:
        df = pd.read_excel(io.BytesIO(fb), sheet_name=s, header=None)
        n  = int(df.notna().sum().sum())
        if n > most: most, best = n, s
    return pd.read_excel(io.BytesIO(fb), sheet_name=best, header=None)

# ══════════════════════════════════════════════════════════════════════
# ATTENDANCE PARSER
# ══════════════════════════════════════════════════════════════════════

def parse_attendance(fb: bytes) -> dict:
    """
    Robust parser for ANY attendance sheet format.
    Finds date row by scanning for the row with the most parseable date values.
    """
    df = best_sheet(fb, hints=["attendance","student"])

    # Find the row with most dates
    best_row, best_cnt, best_first = None, 0, None
    for i, row in df.iterrows():
        hits, first = [], None
        for j, v in enumerate(row):
            d = parse_date(v)
            if d: hits.append(d); first = (j if first is None else first)
        if len(hits) > best_cnt:
            best_cnt, best_row, best_first = len(hits), i, first

    if best_row is None or best_cnt < 2:
        raise ValueError(
            "Could not find session dates. "
            "Ensure dates appear in a single row of your attendance sheet."
        )

    date_row = df.iloc[best_row]
    col_to_date, dates_ordered, seen = {}, [], set()
    for j in range(best_first, df.shape[1]):
        d = parse_date(date_row.iloc[j])
        if d:
            col_to_date[j] = d
            if d not in seen: dates_ordered.append(d); seen.add(d)

    # Find student columns
    name_col = enroll_col = None
    for i in range(best_row):
        row = df.iloc[i]
        rs  = " ".join(str(v).lower() for v in row if not safe_isna(v))
        if any(k in rs for k in ["name","enrol","roll","student"]):
            for j, v in enumerate(row):
                s = str(v).lower().strip()
                if "name" in s and name_col is None:                               name_col   = j
                elif any(k in s for k in ["enrol","roll"]) and enroll_col is None: enroll_col = j
            break

    # Parse students
    students = []
    for i in range(best_row+1, df.shape[0]):
        row = df.iloc[i]
        if row.notna().sum() < 3: continue
        name   = str(row.iloc[name_col]).strip()   if name_col   is not None and not safe_isna(row.iloc[name_col])   else ""
        enroll = str(row.iloc[enroll_col]).strip() if enroll_col is not None and not safe_isna(row.iloc[enroll_col]) else ""
        if not name or name.lower() in ("nan","none",""): continue
        att = {}
        for col_j, d in col_to_date.items():
            val = row.iloc[col_j]
            try:    att[d] = int(float(val)) if not safe_isna(val) else None
            except: att[d] = None
        students.append({"name":name, "enrollment":enroll, "att":att})

    return {"dates": dates_ordered, "students": students}

def group_by_month(dates: list) -> OrderedDict:
    res = defaultdict(list)
    for d in dates:
        if d: res[(d.year, d.month)].append(d)
    return OrderedDict(sorted(res.items()))

# ══════════════════════════════════════════════════════════════════════
# SYLLABUS PARSER — ANY FORMAT
# ══════════════════════════════════════════════════════════════════════

UNIT_RE = re.compile(
    r"^(UNIT|MODULE|CHAPTER|TOPIC|SECTION|PART|BLOCK)\s*[-:.]?\s*\d+",
    re.IGNORECASE
)
SKIP_RE = re.compile(
    r"\bv\.\b|\bAIR\b|\bILR\b|\bSCC\b|https?://|\(\d{4}\)\s+\d+|^\s*Case\s+Law",
    re.IGNORECASE
)

def parse_syllabus(fb: bytes, filename: str) -> OrderedDict:
    ext = filename.lower().rsplit(".", 1)[-1]
    result = OrderedDict()

    def add(unit, text):
        t = text.strip()
        if not t or len(t) < 4 or SKIP_RE.search(t): return
        result.setdefault(unit, []).append(t)

    if ext == "docx":
        import docx as _d
        doc = _d.Document(io.BytesIO(fb)); cur = None
        for p in doc.paragraphs:
            t = p.text.strip()
            if not t: continue
            if UNIT_RE.match(t): cur = t; result[cur] = []
            elif cur: add(cur, t)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if not t: continue
                    if UNIT_RE.match(t): cur = t; result[cur] = []
                    elif cur: add(cur, t)

    elif ext in ("xlsx","xls"):
        df = best_sheet(fb); cur = None
        for _, row in df.iterrows():
            for val in row:
                s = str(val).strip()
                if not s or s.lower() in ("nan","none",""): continue
                if UNIT_RE.match(s): cur = s; result[cur] = []; break
                elif cur and len(s) > 4: add(cur, s); break

    elif ext in ("txt","csv"):
        cur = None
        for line in fb.decode("utf-8","ignore").splitlines():
            t = line.strip()
            if not t: continue
            if UNIT_RE.match(t): cur = t; result[cur] = []
            elif cur: add(cur, t)

    # Fallback: treat whole file as Module 1
    if not result:
        result["Module 1"] = []
        if ext in ("txt","csv"):
            for line in fb.decode("utf-8","ignore").splitlines(): add("Module 1", line)
        elif ext == "docx":
            import docx as _d
            for p in _d.Document(io.BytesIO(fb)).paragraphs: add("Module 1", p.text)
        elif ext in ("xlsx","xls"):
            df = best_sheet(fb)
            for _, row in df.iterrows():
                for val in row:
                    s = str(val).strip()
                    if s and s.lower() not in ("nan","none","") and len(s)>4:
                        add("Module 1", s); break

    return OrderedDict((k,v) for k,v in result.items() if v)

# ══════════════════════════════════════════════════════════════════════
# AUTO-DESCRIPTION
# ══════════════════════════════════════════════════════════════════════

def auto_desc(title: str, module: str) -> str:
    t   = title.strip()
    mod = re.sub(r"^(UNIT|MODULE|CHAPTER)\s*\d+[:\-.]?\s*","",module,flags=re.IGNORECASE).strip() or module
    tl  = t.lower()
    if any(w in tl for w in ["definition","define","meaning"]):
        core = re.sub(r"^(definition\s+(of\s+)?|define\s+)","",t,flags=re.IGNORECASE).strip()
        return (f"This session covers the statutory definition and essential elements of {core}. "
                f"Students will examine the relevant legal provisions, judicial interpretations, "
                f"and practical significance within the framework of {mod}.")
    if any(w in tl for w in ["rights","duties","liability","liabilities","obligation"]):
        return (f"This session examines the rights, duties, and liabilities under {t}. "
                f"Students will analyse statutory provisions, landmark judgments, and the legal "
                f"consequences for parties involved in {mod}.")
    if any(w in tl for w in ["distinction","difference","compare","versus"]):
        return (f"This session provides a comparative analysis of {t}. Students will identify "
                f"key distinctions through statutory provisions and case-law, enabling "
                f"differential reasoning in {mod}.")
    if any(w in tl for w in ["type","kind","classif","categor","nature","form"]):
        return (f"This session explores the types, categories, and conceptual framework of {t}. "
                f"Students will study the legal significance of each category and its "
                f"practical application within {mod}.")
    if any(w in tl for w in ["termination","discharge","revocation","dissolution"]):
        return (f"This session covers the modes of {t} and the legal consequences that follow. "
                f"Students will study statutory provisions, conditions, and judicial "
                f"precedents governing this aspect of {mod}.")
    if any(w in tl for w in ["creation","formation","essential","element"]):
        return (f"This session discusses the formation process and requisite elements of {t}. "
                f"Students will examine statutory requirements, judicial interpretations, "
                f"and practical illustrations relevant to {mod}.")
    if any(w in tl for w in ["remedy","remedies","damages","compensation"]):
        return (f"This session discusses remedies available in {t}. Students will examine "
                f"statutory and equitable remedies, judicial approaches, and "
                f"computation of relief in {mod}.")
    if any(w in tl for w in ["case","judgment"," v."]):
        return (f"This session analyses {t} as a significant judicial decision. Students will "
                f"examine the facts, legal issues, court reasoning, and the precedential "
                f"value of this ruling in {mod}.")
    return (f"This session provides a comprehensive study of {t} within {mod}. "
            f"Students will analyse relevant statutory provisions, judicial precedents, "
            f"and practical applications through structured discussion and case-based learning.")

# ══════════════════════════════════════════════════════════════════════
# TOPIC BALANCER
# ══════════════════════════════════════════════════════════════════════

def balance_topics(unit_configs: list, total_sessions: int) -> list:
    all_items = []
    for uc in unit_configs:
        for t in uc["topics"]:
            all_items.append({"module":uc["module_name"],"tlo":uc["tlo"],"title":t})
    n = len(all_items)
    if n == 0:
        return [{"module":"Session","tlo":"TLO1","title":f"Session {i+1}"}
                for i in range(total_sessions)]
    if n <= total_sessions:
        result = list(all_items)
        for i in range(total_sessions-n):
            base = all_items[i%n]
            result.append({"module":base["module"],"tlo":base["tlo"],
                           "title":base["title"]+" — Revision"})
        return result
    else:
        total_avail = sum(len(uc["topics"]) for uc in unit_configs)
        result = []
        for uc in unit_configs:
            n_take = max(1, round(total_sessions*len(uc["topics"])/total_avail))
            n_take = min(n_take, len(uc["topics"]))
            for t in uc["topics"][:n_take]:
                result.append({"module":uc["module_name"],"tlo":uc["tlo"],"title":t})
        while len(result) < total_sessions:
            result.append(result[-1]|{"title":result[-1]["title"]+" (Cont.)"})
        return result[:total_sessions]

# ══════════════════════════════════════════════════════════════════════
# SESSION SHEET GENERATOR — EXACT DL.xlsx FORMAT
# All values written as plain TEXT strings → no number-parse errors
# ══════════════════════════════════════════════════════════════════════

HEADERS = [
    "Module Name*", "Start Date Time*", "End Date Time*", "Title*",
    "Description*", "Attendance Mandatory*", "TLO",
    "Teaching Faculty Registration ID*"
]
COL_W = [30, 22, 22, 42, 60, 22, 18, 32]

H_FILL = PatternFill("solid", start_color="1F4E79")
H_FONT = Font(bold=True, color="FFFFFF", size=11)
D_FILL = [PatternFill("solid", start_color="D6EAF8"),
          PatternFill("solid", start_color="EBF5FB")]
TN = Side(style="thin", color="BBBBBB")
BD = Border(left=TN, right=TN, top=TN, bottom=TN)
CT = Alignment(horizontal="center", vertical="center", wrap_text=True)
LF = Alignment(horizontal="left",   vertical="center", wrap_text=True)
TEXT_FMT = "@"   # Force cell to be treated as text by Excel / DG platform


def gen_session_sheet(rows: list) -> bytes:
    """
    Generate session sheet with exact DL.xlsx column structure.
    ALL cells explicitly set as TEXT (number_format='@') to prevent
    any field from being interpreted as a number by the DG upload platform.
    """
    wb = Workbook(); ws = wb.active; ws.title = "Session Sheet"
    aligns = [LF, CT, CT, LF, LF, CT, CT, CT]

    # Header row
    for c, (h, w) in enumerate(zip(HEADERS, COL_W), 1):
        cell = ws.cell(1, c, str(h))
        cell.fill   = H_FILL
        cell.font   = H_FONT
        cell.border = BD
        cell.alignment = CT
        cell.number_format = TEXT_FMT          # <-- text format on header too
        ws.column_dimensions[get_column_letter(c)].width = w
    ws.row_dimensions[1].height = 24

    # Data rows
    keys = ["module","start_dt","end_dt","title","description","mandatory","tlo","faculty_id"]
    for i, row in enumerate(rows, 2):
        fill = D_FILL[(i) % 2]
        for c, (key, al) in enumerate(zip(keys, aligns), 1):
            val  = str(row.get(key, ""))        # always a string
            cell = ws.cell(i, c, val)
            cell.fill          = fill
            cell.border        = BD
            cell.alignment     = al
            cell.font          = Font(size=10)
            cell.number_format = TEXT_FMT       # <-- critical: force text
        ws.row_dimensions[i].height = 30

    ws.freeze_panes = "A2"
    buf = io.BytesIO(); wb.save(buf); return buf.getvalue()

# ══════════════════════════════════════════════════════════════════════
# DAY-WISE ATTENDANCE — FIXED MATCHING LOGIC
# ══════════════════════════════════════════════════════════════════════

def gen_daywise_att(tpl_bytes: bytes, att_data: dict, session_date: date) -> bytes:
    """
    FIXED attendance logic:
    - Build enrollment→att dict from master sheet
    - For each student in DG template:
        * If in master and att==1 → PRESENT
        * If in master and att==0 or None → ABSENT
        * If NOT in master → keep PRESENT (don't penalise unknown students)
    """
    wb = load_workbook(io.BytesIO(tpl_bytes)); ws = wb.active

    # Detect column positions from header
    ec = rc = ac = None
    for c in range(1, ws.max_column+1):
        hv = str(ws.cell(1,c).value or "").lower().strip()
        if "email" in hv:                                        ec = c
        elif "registration" in hv or ("reg" in hv and "id" in hv): rc = c
        elif "attendance" in hv:                                 ac = c
    if not all([ec, rc, ac]):
        ec, rc, ac = 1, 2, 3

    # Build master attendance map: enrollment_str → 0/1
    # Default absent only for students that exist in master with 0
    master_att: dict[str, int] = {}
    for st in att_data["students"]:
        enroll = str(st["enrollment"]).strip()
        val    = st["att"].get(session_date)
        if val is None:
            master_att[enroll] = 0       # date exists but no mark → ABSENT
        else:
            try:    master_att[enroll] = int(float(val))
            except: master_att[enroll] = 0

    G_FILL = PatternFill("solid", start_color="C6EFCE")
    R_FILL = PatternFill("solid", start_color="FFC7CE")
    G_FONT = Font(color="006100", bold=True, size=10)
    R_FONT = Font(color="9C0006", bold=True, size=10)
    C_ALGN = Alignment(horizontal="center", vertical="center")

    for r in range(2, ws.max_row+1):
        reg = str(ws.cell(r, rc).value or "").strip()
        if not reg: continue

        # FIXED: only mark ABSENT if reg IS in master and was not present
        if reg in master_att:
            status = "PRESENT" if master_att[reg] == 1 else "ABSENT"
        else:
            status = "PRESENT"   # not in master → keep as PRESENT (no data ≠ absent)

        cell = ws.cell(r, ac)
        cell.value     = status
        cell.fill      = G_FILL if status == "PRESENT" else R_FILL
        cell.font      = G_FONT if status == "PRESENT" else R_FONT
        cell.alignment = C_ALGN

    buf = io.BytesIO(); wb.save(buf); return buf.getvalue()

def build_zip(sess: bytes, dw: dict) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("session_sheet.xlsx", sess)
        for fn, fb in dw.items(): zf.writestr(f"attendance/{fn}", fb)
    buf.seek(0); return buf.getvalue()

# ══════════════════════════════════════════════════════════════════════
#                           MAIN UI
# ══════════════════════════════════════════════════════════════════════

def main():
    st.markdown("""
    <div class="banner">
      <h1>📋 DG Session & Attendance Sheet Generator</h1>
      <div class="sub">Upload any format · Select dates · Configure units · Generate exact DL.xlsx format</div>
      <div class="credit">Created by Dr. Amar Shukla · IILM University, Gurugram</div>
    </div>""", unsafe_allow_html=True)

    for k in ["att","syl","preview","results","_anm","_snm"]:
        if k not in st.session_state: st.session_state[k] = None

    att_data  = st.session_state.att
    syl_data  = st.session_state.syl or OrderedDict()
    all_dates = att_data["dates"] if att_data else []
    month_map = group_by_month(all_dates)
    syl_units = list(syl_data.keys())

    # Stats strip
    if all_dates or syl_data:
        n_st = len(att_data["students"]) if att_data else 0
        n_tp = sum(len(v) for v in syl_data.values())
        st.markdown(f"""
        <div class="stat-row">
          <div class="stat-card"><div class="stat-num">{len(all_dates)}</div>
              <div class="stat-lbl">Session Dates</div></div>
          <div class="stat-card"><div class="stat-num">{len(month_map)}</div>
              <div class="stat-lbl">Months Available</div></div>
          <div class="stat-card"><div class="stat-num">{n_st}</div>
              <div class="stat-lbl">Students</div></div>
          <div class="stat-card"><div class="stat-num">{len(syl_units)}</div>
              <div class="stat-lbl">Units in Syllabus</div></div>
          <div class="stat-card"><div class="stat-num">{n_tp}</div>
              <div class="stat-lbl">Topics Found</div></div>
        </div>""", unsafe_allow_html=True)

    st.divider()

    # ══════════════════════════════════════════════════════════════════
    # STEP 1 — UPLOADS
    # ══════════════════════════════════════════════════════════════════
    sec("📂  Step 1 — Upload Files")

    uc1, uc2, uc3 = st.columns(3)
    with uc1:
        st.markdown("**📅 Master Attendance Sheet**")
        st.caption("Any Excel format — dates are auto-detected regardless of layout")
        att_file = st.file_uploader("",type=["xlsx","xls"],key="u_att",label_visibility="collapsed")
        if att_file: st.markdown(f'<span class="badge">✅ {att_file.name}</span>',unsafe_allow_html=True)

    with uc2:
        st.markdown("**📚 Syllabus File**")
        st.caption("Accepts .docx · .xlsx · .txt · .csv — units detected automatically")
        syl_file = st.file_uploader("",type=["docx","xlsx","xls","txt","csv"],key="u_syl",label_visibility="collapsed")
        if syl_file: st.markdown(f'<span class="badge">✅ {syl_file.name}</span>',unsafe_allow_html=True)

    with uc3:
        st.markdown("**🗂️ DG Attendance Template**")
        st.caption("The pre-filled template with student Email IDs and Registration IDs")
        att_tpl = st.file_uploader("",type=["xlsx"],key="u_atpl",label_visibility="collapsed")
        if att_tpl: st.markdown(f'<span class="badge">✅ {att_tpl.name}</span>',unsafe_allow_html=True)

    # Auto-parse files
    if att_file and att_file.name != st.session_state._anm:
        try:
            st.session_state.att  = parse_attendance(att_file.read())
            st.session_state._anm = att_file.name
            st.session_state.preview = None; st.session_state.results = None
            att_data  = st.session_state.att
            all_dates = att_data["dates"]
            month_map = group_by_month(all_dates)
            al_ok(f"Attendance loaded — <b>{len(all_dates)} session dates</b> · "
                  f"<b>{len(att_data['students'])} students</b>")
        except Exception as ex: al_err(f"Could not read attendance: {ex}")

    if syl_file and syl_file.name != st.session_state._snm:
        try:
            st.session_state.syl  = parse_syllabus(syl_file.read(), syl_file.name)
            st.session_state._snm = syl_file.name
            syl_data  = st.session_state.syl
            syl_units = list(syl_data.keys())
            n_tp      = sum(len(v) for v in syl_data.values())
            al_ok(f"Syllabus loaded — <b>{len(syl_units)} unit(s)</b> · <b>{n_tp} topics</b>")
        except Exception as ex: al_warn(f"Syllabus warning: {ex}")

    st.divider()

    # ══════════════════════════════════════════════════════════════════
    # STEP 2 — SELECT DATES (month → individual dates)
    # ══════════════════════════════════════════════════════════════════
    sec("📅  Step 2 — Select Session Dates")

    if not month_map:
        al_info("Upload the Master Attendance Sheet first — months and dates will appear here.")
        selected_dates = []
    else:
        al_info(
            "Tick the months to include. "
            "Expand each month to <b>select individual dates</b> — "
            "all dates are pre-selected, uncheck any date to exclude it."
        )
        selected_dates = []

        for row_start in range(0, len(month_map), 4):
            keys_row = list(month_map.keys())[row_start:row_start+4]
            cols     = st.columns(len(keys_row))

            for ci, ym in enumerate(keys_row):
                ds    = month_map[ym]
                label = datetime(ym[0],ym[1],1).strftime("%B %Y")
                with cols[ci]:
                    month_on = st.checkbox(f"**{label}**", value=True,
                                           key=f"m_{ym[0]}_{ym[1]}")
                    st.caption(f"🗓️ {len(ds)} session dates")

                    if month_on:
                        with st.expander(f"Choose dates in {label}", expanded=False):
                            al_info("Uncheck any date to exclude it from session generation.")
                            for d in ds:
                                d_lbl = f"{d.strftime('%d %b %Y')}  ({d.strftime('%A')})"
                                if st.checkbox(d_lbl, value=True,
                                               key=f"d_{d.isoformat()}"):
                                    selected_dates.append(d)
                        # Show chip preview
                        chips = "".join(
                            f'<span class="date-chip">{d.strftime("%d")} '
                            f'{d.strftime("%a")}</span>' for d in ds
                        )
                        st.markdown(chips, unsafe_allow_html=True)

        selected_dates = sorted(set(selected_dates))
        if selected_dates:
            al_ok(f"<b>{len(selected_dates)} dates selected</b> — "
                  f"{selected_dates[0].strftime('%d %b %Y')} to "
                  f"{selected_dates[-1].strftime('%d %b %Y')}")
        else:
            al_warn("Select at least one date to proceed.")

    st.divider()

    # ══════════════════════════════════════════════════════════════════
    # STEP 3 — SESSION COUNT + BASIC INFO
    # ══════════════════════════════════════════════════════════════════
    sec("⚙️  Step 3 — How Many Sessions to Generate?")

    al_info(
        "<b>Sessions to generate</b> = number of rows in the output session sheet. "
        "By default this equals the number of selected dates. "
        "You can enter a custom number — e.g. if you selected 30 dates "
        "but want only 25 sessions, type 25."
    )

    sc = st.columns([1.5, 1.5, 1, 1, 1])

    with sc[0]:
        st.markdown("**📊 Total Sessions to Generate**")
        custom_count = st.checkbox(
            f"Enter a custom number  "
            f"(default = {len(selected_dates)} from selected dates)",
            value=False, key="cust"
        )
        if custom_count:
            total_sessions = st.number_input(
                "Number of sessions:", min_value=1, max_value=500,
                value=len(selected_dates) or 50
            )
        else:
            total_sessions = len(selected_dates)
            st.metric("Sessions (matches selected dates)", total_sessions)

    with sc[1]:
        faculty_id = st.text_input(
            "🧑‍🏫 Faculty Registration ID",
            placeholder="e.g. IILMGG006412025",
            help="Appears in every row of the session sheet — 'Teaching Faculty Registration ID*'"
        )

    with sc[2]:
        mandatory = st.selectbox("✅ Attendance Mandatory?", ["TRUE","FALSE"])

    with sc[3]:
        tlo_max = st.number_input("Max TLO Number", min_value=1, max_value=100, value=5)

    with sc[4]:
        tlo_prefix = st.text_input("TLO Prefix", value="TLO",
                                   help="'TLO' → TLO1, TLO2 | 'CO' → CO1, CO2")

    all_tlos = [f"{tlo_prefix}{i}" for i in range(1, tlo_max+1)]

    st.divider()

    # ══════════════════════════════════════════════════════════════════
    # STEP 4 — TIMING
    # ══════════════════════════════════════════════════════════════════
    sec("⏰  Step 4 — Lecture Timing (Day-wise)")
    al_info(
        "Enable the days on which lectures are held. "
        "Set Start and End time for each enabled day. "
        "Days detected in your selected dates are pre-enabled."
    )

    detected = {}
    for d in selected_dates:
        dn = d.strftime("%A"); detected[dn] = detected.get(dn,0)+1

    day_timing = {}
    cols7 = st.columns(7)

    for i, day in enumerate(ALL_DAYS):
        cnt = detected.get(day, 0)
        with cols7[i]:
            enabled = st.checkbox(f"**{SHORT[day]}**", value=(cnt>0), key=f"de_{day}")
            st.caption(f"🔵 {cnt} sessions" if cnt else "—")

            if enabled:
                h_list = list(range(1,13))
                m_list = ["00","05","10","15","20","25","30","35","40","45","50","55"]
                h1,m1,a1 = st.columns(3)
                with h1: sh=st.selectbox("",h_list, index=9, key=f"sh_{day}",label_visibility="collapsed")
                with m1: sm=st.selectbox("",m_list, index=0, key=f"sm_{day}",label_visibility="collapsed")
                with a1: sa=st.selectbox("",["AM","PM"],index=1,key=f"sa_{day}",label_visibility="collapsed")
                s24=to_24h(sh,sm,sa)
                h2,m2,a2 = st.columns(3)
                with h2: eh=st.selectbox("",h_list, index=10,key=f"eh_{day}",label_visibility="collapsed")
                with m2: em=st.selectbox("",m_list, index=0, key=f"em_{day}",label_visibility="collapsed")
                with a2: ea=st.selectbox("",["AM","PM"],index=1,key=f"ea_{day}",label_visibility="collapsed")
                e24=to_24h(eh,em,ea)
                day_timing[day]=(s24,e24)
                st.markdown(f'<div class="time-tag">⏱ {s24}–{e24}</div>', unsafe_allow_html=True)

    st.divider()

    # ══════════════════════════════════════════════════════════════════
    # STEP 5 — HOW MANY UNITS + MODULE CONFIG
    # ══════════════════════════════════════════════════════════════════
    sec("📚  Step 5 — Units to Include & Module Configuration")

    unit_configs = []

    if not syl_units:
        al_warn(
            "No syllabus uploaded. Session titles will default to 'Session 1', 'Session 2', etc. "
            "Upload a syllabus to auto-fill titles and descriptions."
        )
    else:
        al_info(
            f"<b>{len(syl_units)} unit(s)</b> detected in your syllabus. "
            "Choose <b>how many units to include</b> in this session sheet. "
            "For each included unit, enter a Module Name, assign TLOs, "
            "and choose Full Unit or Custom topic selection."
        )

        # HOW MANY UNITS TO INCLUDE — key question
        n_include = st.number_input(
            f"How many units to include?  (Syllabus has {len(syl_units)} units)",
            min_value=1,
            max_value=len(syl_units),
            value=len(syl_units),
            help=(
                "Example: Syllabus has 4 units but you only want to cover "
                "2 units in this session sheet — enter 2."
            )
        )

        units_to_use = syl_units[:n_include]

        if n_include < len(syl_units):
            al_info(
                f"Including <b>{n_include} of {len(syl_units)} units</b>: "
                + ", ".join(f"Unit {i+1}" for i in range(n_include))
            )

        for i, unit_key in enumerate(units_to_use):
            avail = syl_data[unit_key]

            with st.expander(
                f"Unit {i+1}  ·  {unit_key[:65]}  ({len(avail)} topics available)",
                expanded=(i < 3)
            ):
                r1, r2 = st.columns([3, 2])

                with r1:
                    mod_name = st.text_input(
                        "Module Name*  — this text appears in the 'Module Name*' column",
                        value=unit_key,
                        key=f"mn_{i}",
                        help="Edit this to set exactly what appears in the session sheet output."
                    )

                with r2:
                    def_tlo  = [all_tlos[i % len(all_tlos)]] if all_tlos else [f"{tlo_prefix}1"]
                    sel_tlos = st.multiselect(
                        "TLOs for this module",
                        options=all_tlos,
                        default=def_tlo,
                        key=f"tl_{i}",
                        help="Multiple TLOs will appear as 'TLO1 | TLO2' in the sheet."
                    )
                    tlo_str  = " | ".join(sel_tlos) if sel_tlos else all_tlos[i%len(all_tlos)]

                cov = st.radio(
                    "Topic Coverage for this module:",
                    ["📖 Full Unit — include all topics",
                     "✂️ Custom — I will select specific topics"],
                    key=f"cov_{i}", horizontal=True
                )

                if "Custom" in cov:
                    chosen = st.multiselect(
                        "Select topics to include:",
                        options=avail, default=avail, key=f"tp_{i}",
                        help="All topics pre-selected — remove any you don't want."
                    )
                    final_topics = chosen
                    st.caption(f"✅ {len(final_topics)} of {len(avail)} topics selected")
                else:
                    final_topics = avail
                    with st.expander(f"Preview all {len(avail)} topics →"):
                        for j,t in enumerate(avail,1): st.write(f"`{j}.` {t}")

                unit_configs.append({
                    "module_name": mod_name,
                    "tlo":         tlo_str,
                    "topics":      final_topics,
                })

        # Summary
        n_topics = sum(len(u["topics"]) for u in unit_configs)
        sc2 = st.columns(3)
        with sc2[0]:
            st.markdown(
                f'<div class="stat-card"><div class="stat-num">{len(unit_configs)}</div>'
                f'<div class="stat-lbl">Units Included</div></div>',
                unsafe_allow_html=True)
        with sc2[1]:
            st.markdown(
                f'<div class="stat-card"><div class="stat-num">{n_topics}</div>'
                f'<div class="stat-lbl">Topics Selected</div></div>',
                unsafe_allow_html=True)
        with sc2[2]:
            st.markdown(
                f'<div class="stat-card"><div class="stat-num">{total_sessions}</div>'
                f'<div class="stat-lbl">Sessions to Generate</div></div>',
                unsafe_allow_html=True)

        st.markdown("<br>",unsafe_allow_html=True)
        if n_topics == total_sessions:
            al_ok("Perfect match — one topic per session.")
        elif n_topics < total_sessions:
            al_warn(f"Topics ({n_topics}) < Sessions ({total_sessions}). "
                    f"Last {total_sessions-n_topics} sessions will be revision entries.")
        else:
            al_warn(f"Topics ({n_topics}) > Sessions ({total_sessions}). "
                    f"Topics will be trimmed proportionally.")

    st.divider()

    # ══════════════════════════════════════════════════════════════════
    # STEP 6 — BUILD PREVIEW
    # ══════════════════════════════════════════════════════════════════
    sec("👀  Step 6 — Preview & Edit Table")
    al_info(
        "Click <b>Build Preview Table</b> to generate all rows. "
        "Topics are auto-balanced across modules and descriptions are auto-generated. "
        "Every cell is editable in the table — make any changes before generating files."
    )

    prev_btn = st.button("🔄  Build Preview Table", type="secondary",
                         use_container_width=True)

    if prev_btn:
        if not selected_dates:
            al_err("Select at least one date in Step 2.")
        elif not faculty_id.strip():
            al_err("Enter Faculty Registration ID in Step 3.")
        else:
            flat = (balance_topics(unit_configs, total_sessions) if unit_configs
                    else [{"module":"Module","tlo":f"{tlo_prefix}1","title":f"Session {i+1}"}
                          for i in range(total_sessions)])

            rows = []
            for idx in range(total_sessions):
                if idx < len(selected_dates):
                    d   = selected_dates[idx]
                    ds  = d.strftime("%Y-%m-%d")
                    dn  = d.strftime("%A")
                    s,e = day_timing.get(dn, ("10:00","11:00"))
                    sdt = f"{ds} {s}:00"
                    edt = f"{ds} {e}:00"
                else:
                    sdt = edt = "TBD"

                sess = (flat[idx] if idx < len(flat) else
                        {"module":"Extra","tlo":f"{tlo_prefix}1",
                         "title":f"Extra Session {idx+1}"})

                rows.append({
                    "Sr":               idx+1,
                    "Module Name*":     sess["module"],
                    "Start Date Time*": sdt,
                    "End Date Time*":   edt,
                    "Title*":           sess["title"],
                    "Description*":     auto_desc(sess["title"], sess["module"]),
                    "Mandatory*":       mandatory,
                    "TLO":              sess["tlo"],
                    "Faculty Reg ID*":  faculty_id.strip(),
                })

            st.session_state.preview = pd.DataFrame(rows)
            st.session_state.results = None
            al_ok(f"Preview ready — <b>{len(rows)} rows</b>. Edit the table below if needed.")

    if st.session_state.preview is not None:
        edited = st.data_editor(
            st.session_state.preview,
            use_container_width=True, num_rows="fixed",
            hide_index=True, key="ed_prev",
            column_config={
                "Sr":               st.column_config.NumberColumn("Sr",width=55,disabled=True),
                "Module Name*":     st.column_config.TextColumn("Module Name* ✏️",    width=235),
                "Start Date Time*": st.column_config.TextColumn("Start DateTime* ✏️", width=170),
                "End Date Time*":   st.column_config.TextColumn("End DateTime* ✏️",   width=170),
                "Title*":           st.column_config.TextColumn("Title* ✏️",          width=265),
                "Description*":     st.column_config.TextColumn("Description* ✏️",    width=385),
                "Mandatory*":       st.column_config.SelectboxColumn(
                                        "Mandatory*",options=["TRUE","FALSE"],width=115),
                "TLO":              st.column_config.TextColumn("TLO ✏️",             width=130),
                "Faculty Reg ID*":  st.column_config.TextColumn("Faculty Reg ID* ✏️", width=170),
            },
            height=500,
        )
        st.session_state.preview = edited

    st.divider()

    # ══════════════════════════════════════════════════════════════════
    # STEP 7 — GENERATE
    # ══════════════════════════════════════════════════════════════════
    sec("🚀  Step 7 — Generate Files")

    can_gen = bool(
        st.session_state.preview is not None and
        att_data and faculty_id.strip()
    )
    if not can_gen:
        miss = []
        if not att_data:                       miss.append("Attendance Sheet (Step 1)")
        if not faculty_id.strip():             miss.append("Faculty Registration ID (Step 3)")
        if st.session_state.preview is None:  miss.append("Preview Table (Step 6)")
        if miss: al_info(f"Complete first: <b>{' · '.join(miss)}</b>")

    gen_btn = st.button(
        "⚡  Generate Session Sheet + Day-wise Attendance Files",
        type="primary", disabled=not can_gen, use_container_width=True
    )

    if gen_btn and can_gen:
        errors=[]; prog=st.progress(0); log=st.empty()
        try:
            df = st.session_state.preview
            log.markdown('<div class="al-info">📊 Building session rows…</div>',
                         unsafe_allow_html=True)
            session_rows = [{
                "module":      str(r.get("Module Name*","")),
                "start_dt":    str(r.get("Start Date Time*","")),
                "end_dt":      str(r.get("End Date Time*","")),
                "title":       str(r.get("Title*","")),
                "description": str(r.get("Description*","")),
                "mandatory":   str(r.get("Mandatory*","TRUE")),
                "tlo":         str(r.get("TLO",f"{tlo_prefix}1")),
                "faculty_id":  str(r.get("Faculty Reg ID*", faculty_id.strip())),
            } for _,r in df.iterrows()]
            prog.progress(20)

            log.markdown('<div class="al-info">📊 Generating session sheet…</div>',
                         unsafe_allow_html=True)
            session_xlsx = gen_session_sheet(session_rows)
            al_ok(f"Session sheet ready — <b>{len(session_rows)} rows</b> "
                  f"in exact DL.xlsx format. All cells are plain text — no format errors.")
            prog.progress(50)

            daywise = {}
            if att_tpl:
                log.markdown('<div class="al-info">🗂️ Generating day-wise attendance…</div>',
                             unsafe_allow_html=True)
                tpl_b = att_tpl.read()
                bar   = st.progress(0)
                use_d = selected_dates[:total_sessions]
                for i,d in enumerate(use_d):
                    try:
                        fn = f"attendance_{d.strftime('%Y-%m-%d')}.xlsx"
                        daywise[fn] = gen_daywise_att(tpl_b, att_data, d)
                    except Exception as ex: errors.append(f"{d}: {ex}")
                    bar.progress((i+1)/max(len(use_d),1))
                al_ok(f"Day-wise attendance — <b>{len(daywise)} files</b> ready. "
                      f"PRESENT/ABSENT correctly matched from master attendance.")
            else:
                al_warn("No DG Attendance Template uploaded — only session sheet available.")

            prog.progress(90)
            zip_b = build_zip(session_xlsx, daywise)
            prog.progress(100)
            log.markdown(
                '<div class="al-ok">🎉 All files generated successfully!</div>',
                unsafe_allow_html=True)
            st.session_state.results = {
                "session":session_xlsx,"daywise":daywise,
                "zip":zip_b,"errors":errors,"n":len(session_rows)
            }
        except Exception as ex:
            al_err(f"Error: {ex}")
            import traceback; st.code(traceback.format_exc())

    # ══════════════════════════════════════════════════════════════════
    # STEP 8 — DOWNLOADS
    # ══════════════════════════════════════════════════════════════════
    if st.session_state.results:
        res = st.session_state.results
        st.divider()
        sec("📥  Step 8 — Download")
        for e in res.get("errors",[]): al_warn(f"Skipped: {e}")

        d1,d2 = st.columns(2)
        with d1:
            st.download_button(
                "📦  Download Everything as ZIP",
                data=res["zip"],
                file_name=f"DG_Output_{datetime.now():%Y%m%d_%H%M}.zip",
                mime="application/zip",
                use_container_width=True
            )
        with d2:
            st.download_button(
                "📋  Download session_sheet.xlsx only",
                data=res["session"], file_name="session_sheet.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )

        if res["daywise"]:
            st.markdown(f"#### 📅  Day-wise Attendance  ({len(res['daywise'])} files)")
            files = list(res["daywise"].items())
            for ri in range(0,len(files),4):
                chunk=files[ri:ri+4]; cols=st.columns(4)
                for ci,(fn,fb) in enumerate(chunk):
                    with cols[ci]:
                        dp=fn.replace("attendance_","").replace(".xlsx","")
                        try:    lbl=datetime.strptime(dp,"%Y-%m-%d").strftime("%d %b %Y")
                        except: lbl=dp
                        st.download_button(f"📅  {lbl}",data=fb,file_name=fn,
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True,key=f"dl_{fn}")

        st.markdown("<br>",unsafe_allow_html=True)
        if st.button("🔄  Reset — Start Over"):
            for k in ["att","syl","preview","results","_anm","_snm"]:
                st.session_state[k]=None
            st.rerun()

    st.markdown("""
    <div class="footer">
      Created by <b>Dr. Amar Shukla</b> · IILM University, Gurugram ·
      All processing is in-memory — no data is stored on any server
    </div>""", unsafe_allow_html=True)


if __name__ == "__main__":
    main()
